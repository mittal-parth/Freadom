from django.shortcuts import render, HttpResponse
from django.http import JsonResponse
import requests
import sys
import os
from docx import Document
from datetime import datetime
from subprocess import run, PIPE

# Create your views here.
def index(request):
    if request.method == 'POST':
        url = request.POST.get('url')
        dirname = os.path.dirname(__file__)
        script = os.path.join(dirname,'script.py')
        output = run([sys.executable,script,url], shell=False, stdout=PIPE)
        # print(output.stdout.decode('unicode_escape'))
        output_filename = 'webscraper_result' + f"{datetime.now().strftime('%D %M %Y %H %M %S')}" + '.docx'
        output_file_path = os.getcwd() + '\\media\\' + output_filename

        # mydoc = docx.Document()
        # mydoc.add_paragraph(output.stdout.decode('unicode_escape'))
        # print("Para added")
        # mydoc.save(output_file_path)    

        print("Yupppp")
        with open(output_file_path, 'w') as f:
            print("inside function")
            source_stream = StringIO(f.read())
            document = Document(source_stream)
            source_stream.close()
            target_stream = StringIO()
            document.save(target_stream)

        if os.path.exists(output_file_path):
            with open(output_file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=web_scraper_result.docx'
                return response
        else:
            return HttpResponse('Oops! Something went wrong.')
    return render(request, 'index.html')